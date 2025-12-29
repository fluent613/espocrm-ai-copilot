# resume_parser.py
# Resume parsing functionality for FluencyCare Copilot

import re
import json
import logging
from typing import Tuple, Optional, Dict, Any
import openai
from utils import format_phone_for_crm, create_phone_number_data

logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self, openai_client):
        self.client = openai_client

    def extract_name_from_filename(self, filename: str) -> Optional[Tuple[str, str]]:
        """Extract name from filename like 'John_Doe_Resume.pdf' or 'Jane Smith.docx'"""
        if not filename:
            return None

        # Remove file extension
        name_part = re.sub(r'\.(pdf|docx?|txt)$', '', filename, flags=re.IGNORECASE)

        # Remove common resume-related words
        name_part = re.sub(r'\b(resume|cv|curriculum|vitae)\b', '', name_part, flags=re.IGNORECASE)

        # Replace underscores and hyphens with spaces
        name_part = name_part.replace('_', ' ').replace('-', ' ')

        # Clean up extra spaces
        name_part = re.sub(r'\s+', ' ', name_part).strip()

        # Try to match "FirstName LastName" pattern
        name_match = re.match(r'^([A-Z][a-z]+)\s+([A-Z][a-z]+)', name_part, re.IGNORECASE)
        if name_match:
            first = name_match.group(1).capitalize()
            last = name_match.group(2).capitalize()
            logger.info(f"Extracted name from filename '{filename}': {first} {last}")
            return first, last

        logger.warning(f"Could not extract name from filename: {filename}")
        return None

    def manual_name_extraction(self, resume_text: str) -> Optional[Tuple[str, str]]:
        """Manual name extraction with better patterns and validation"""
        # Common words that are NOT person names (section headers, job titles, etc.)
        INVALID_NAMES = {
            'machine', 'learning', 'deep', 'data', 'science', 'engineer', 'developer',
            'summary', 'experience', 'education', 'skills', 'professional', 'technical',
            'areas', 'expertise', 'certifications', 'honors', 'awards', 'publications',
            'contact', 'resume', 'curriculum', 'vitae', 'profile', 'objective',
            'qualifications', 'background', 'overview', 'highlights', 'analyst',
            'scientist', 'specialist', 'manager', 'director', 'senior', 'junior',
            'software', 'hardware', 'systems', 'network', 'database', 'cloud',
            'predictive', 'analytics', 'visualization', 'intelligence', 'artificial',
            'natural', 'language', 'processing', 'computer', 'vision',
            # CRITICAL: Add section headers that are commonly mistaken for names
            'work', 'history', 'career', 'employment', 'positions', 'responsibilities',
            'accomplishments', 'achievements', 'projects', 'activities', 'interests',
            'references', 'certifications', 'licenses', 'training', 'courses',
            'volunteer', 'leadership', 'memberships', 'presentations', 'patents'
        }

        lines = resume_text.strip().split('\n')

        # Check first 10 lines for name patterns
        for line in lines[:10]:
            line = line.strip()
            if not line or len(line) < 3:
                continue

            # Remove common prefixes/suffixes and clean line
            clean_line = re.sub(r'\b(Mr\.?|Ms\.?|Mrs\.?|Dr\.?)\s+', '', line, flags=re.IGNORECASE)
            clean_line = re.sub(r'\s+(Jr\.?|Sr\.?|II|III)$', '', clean_line, flags=re.IGNORECASE)
            clean_line = clean_line.strip()

            # Look for simple "FirstName LastName" pattern (2-20 chars each)
            name_match = re.match(r'^([A-Z][a-z]{1,20})\s+([A-Z][a-z]{1,20})$', clean_line)
            if name_match:
                first, last = name_match.group(1), name_match.group(2)
                # Validate it's not a skill/section name
                if first.lower() not in INVALID_NAMES and last.lower() not in INVALID_NAMES:
                    logger.info(f"Manual extraction validated: {first} {last}")
                    return first, last

            # Look for "FIRSTNAME LASTNAME" pattern
            caps_match = re.match(r'^([A-Z]{2,20})\s+([A-Z]{2,20})$', clean_line)
            if caps_match:
                first, last = caps_match.group(1).capitalize(), caps_match.group(2).capitalize()
                if first.lower() not in INVALID_NAMES and last.lower() not in INVALID_NAMES:
                    logger.info(f"Manual extraction validated (caps): {first} {last}")
                    return first, last

            # Look for "First Last" with possible middle initial
            middle_match = re.match(r'^([A-Z][a-z]{1,20})\s+[A-Z]\.?\s+([A-Z][a-z]{1,20})$', clean_line)
            if middle_match:
                first, last = middle_match.group(1), middle_match.group(2)
                if first.lower() not in INVALID_NAMES and last.lower() not in INVALID_NAMES:
                    logger.info(f"Manual extraction validated (with middle): {first} {last}")
                    return first, last

        logger.warning("Manual extraction found no valid person names")
        return None

    def extract_resume_info(self, resume_text: str, filename: str = None) -> Dict[str, Any]:
        """Extract structured information from resume text with improved parsing"""
        prompt = f"""
        Please extract the following information from this resume and return it as JSON:
        - firstName: First name only
        - lastName: Last name only (REQUIRED - if you cannot find a clear last name, use "Unknown" or extract from email)
        - emailAddress: Email address (must contain @ symbol)
        - phoneNumber: Phone number (10+ digits, format: (XXX) XXX-XXXX or similar)
        - cCurrentTitle: Most recent job title or current position
        - cSkills: Array of relevant technical skills and technologies
        - cCurrentCompany: Most recent company name
        - cLinkedInURL: LinkedIn profile URL if found
        - addressStreet: Street address (number and street name)
        - addressCity: City name
        - addressState: State (2-letter code or full name)
        - addressPostalCode: ZIP/Postal code
        - addressCountry: Country (default to USA for US addresses)

        CRITICAL PARSING RULES FOR NAMES (FOLLOW STRICTLY):
        - The person's name is ALWAYS in the first 1-3 lines of the resume - NEVER use section headers as names!
        - Look at the VERY FIRST LINE FIRST - this is almost always the person's name
        - Common name patterns: "John Smith", "JOHN SMITH", "John A. Smith", "Smith, John"
        - INVALID NAMES (these are section headers, NOT names): "Work History", "Professional Experience", "Summary", "Objective", "Education", "Skills", "Career Summary", "Professional Profile"
        - If the first line is a section header, check the SECOND line for the name
        - Names are typically 2-4 words, title case or all caps, with no special characters except periods or hyphens
        - If you extract a name and it seems like a section header, YOU ARE WRONG - look again at the first 1-2 lines
        - Example: If you see "Work History" in line 5 but "Kamlesh Dhairyavan" in line 1, the name is "Kamlesh" "Dhairyavan"

        CRITICAL PHONE NUMBER RULES:
        - Phone numbers can appear in many formats: (303) 555-1234, 303-555-1234, 303.555.1234, 303 555 1234
        - Look in the header section (first 5-10 lines) for phone numbers
        - Common patterns:
          * "NAME | PHONE | EMAIL" (pipe-separated header)
          * "Phone: (XXX) XXX-XXXX"
          * "Mobile: XXX-XXX-XXXX"
          * "Tel: XXX.XXX.XXXX"
        - Phone numbers are ALWAYS 10 digits in US format
        - DO NOT confuse with ZIP codes or other numbers

        EMAIL EXTRACTION:
        - Look in the header (first 5-10 lines)
        - Format: anything@domain.com
        - Often appears alongside phone in header: "NAME | PHONE | EMAIL"

        ADDRESS vs PHONE RULES:
        - Phone numbers: 10 digits in XXX-XXX-XXXX pattern
        - Street addresses: house number + street name (like "123 Main Street")
        - City names: words, not number patterns
        - Never put phone numbers in address fields!

        Resume text: {resume_text}

        Return only valid JSON with these exact field names. REMEMBER: The person's name is in the first 1-3 lines, NOT in section headers!
        """
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4.1",
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"},
                timeout=30
            )
            
            result = json.loads(response.choices[0].message.content)
            logger.info(f"OpenAI extracted: {result}")
            logger.info(f"🔍 AI EXTRACTED NAME: {result.get('firstName', 'N/A')} {result.get('lastName', 'N/A')}")

            # Post-process and validate the results
            cleaned_result = {}
            
            # Clean and validate names
            firstName = str(result.get('firstName', '')).strip() if result.get('firstName') else ""
            lastName = str(result.get('lastName', '')).strip() if result.get('lastName') else ""
            
            # If AI didn't find good names, try manual extraction first
            if (not firstName or not lastName or
                firstName.lower() in ['unknown', 'resume', 'contact'] or
                lastName.lower() in ['professional', 'unknown', 'contact']):

                logger.info("AI name extraction poor, trying manual extraction...")
                manual_names = self.manual_name_extraction(resume_text)
                if manual_names:
                    firstName, lastName = manual_names
                    logger.info(f"Manual extraction found: {firstName} {lastName}")
                else:
                    # Try extracting from filename as last resort
                    logger.info("Manual extraction failed, trying filename extraction...")
                    if filename:
                        filename_names = self.extract_name_from_filename(filename)
                        if filename_names:
                            firstName, lastName = filename_names
                            logger.info(f"Filename extraction found: {firstName} {lastName}")
            
            # If still no lastName, try to extract from email
            if not lastName and result.get('emailAddress'):
                email = str(result['emailAddress'])
                if '.' in email.split('@')[0]:  # Has dot in username part
                    email_parts = email.split('@')[0].split('.')
                    if len(email_parts) >= 2:
                        if not firstName:
                            firstName = email_parts[0].capitalize()
                        lastName = email_parts[-1].capitalize()
                        logger.info(f"Extracted names from email: {firstName} {lastName}")
            
            # ONLY use these fallbacks if everything else completely fails
            if not firstName:
                firstName = "Contact"
            if not lastName:
                lastName = "Person"
            
            cleaned_result['firstName'] = firstName
            cleaned_result['lastName'] = lastName
            logger.info(f"✅ FINAL EXTRACTED NAME: {firstName} {lastName}")

            # Validate and clean other fields
            if result.get('emailAddress') and '@' in str(result['emailAddress']):
                cleaned_result['emailAddress'] = str(result['emailAddress']).strip()
                
            if result.get('phoneNumber'):
                cleaned_result['phoneNumber'] = str(result['phoneNumber']).strip()
                
            # Clean other fields
            for field in ['cCurrentTitle', 'cCurrentCompany', 'cLinkedInURL', 
                         'addressStreet', 'addressCity', 'addressState', 
                         'addressPostalCode', 'addressCountry']:
                if result.get(field):
                    cleaned_result[field] = str(result[field]).strip()
            
            # Handle skills (can be array or string)
            if result.get('cSkills'):
                if isinstance(result['cSkills'], list):
                    cleaned_result['cSkills'] = ', '.join(result['cSkills'])
                else:
                    cleaned_result['cSkills'] = str(result['cSkills']).strip()
                    
            logger.info(f"Cleaned result: {cleaned_result}")
            return cleaned_result
            
        except Exception as e:
            logger.error(f"Resume parsing error: {e}")
            return self._fallback_parsing(resume_text)

    def _fallback_parsing(self, resume_text: str) -> Dict[str, Any]:
        """Enhanced fallback parsing using regex"""
        # Clean the text first - remove garbled characters
        clean_text = re.sub(r'[^\x00-\x7F]+', ' ', resume_text)
        clean_text = re.sub(r'\s+', ' ', clean_text)
        
        # Try manual extraction first
        manual_names = self.manual_name_extraction(resume_text)
        if manual_names:
            firstName, lastName = manual_names
        else:
            firstName = "Contact"
            lastName = "Person"
        
        # Find email
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', clean_text)
        
        # Find phone number
        phone_match = re.search(r'\(?\d{3}\)?[\s\-\.]?\d{3}[\s\-\.]?\d{4}', clean_text)
        
        # Find address components
        address_match = re.search(r'(\d+\s+[^,\n]+),?\s*([^,\n]+),?\s*([A-Z]{2})\s+(\d{5}(?:-\d{4})?)', clean_text)
        
        result = {
            "firstName": firstName,
            "lastName": lastName,
            "emailAddress": email_match.group(0) if email_match else "",
            "phoneNumber": phone_match.group(0) if phone_match else "",
            "cSkills": "",
            "cCurrentTitle": "",
            "cCurrentCompany": "",
            "cLinkedInURL": ""
        }
        
        # Add address components if found
        if address_match:
            result.update({
                "addressStreet": address_match.group(1).strip(),
                "addressCity": address_match.group(2).strip(),
                "addressState": address_match.group(3).strip(),
                "addressPostalCode": address_match.group(4).strip(),
                "addressCountry": "USA"
            })
        else:
            result.update({
                "addressStreet": "",
                "addressCity": "",
                "addressState": "",
                "addressPostalCode": "",
                "addressCountry": ""
            })
        
        return result

    def process_uploaded_file(self, file) -> Tuple[Optional[str], Optional[str]]:
        """Process uploaded file and extract text content with better error handling"""
        try:
            # Validate file object
            if not file:
                logger.error("No file object provided")
                return None, "❌ No file provided. Please select a file to upload."

            filename = file.filename
            if not filename:
                logger.error("File has no filename")
                return None, "❌ Invalid file. Please select a valid file to upload."

            filename_lower = filename.lower()
            logger.info(f"Processing file: {filename}")

            # Initialize content variable
            content = None

            if filename_lower.endswith('.txt'):
                content = file.read().decode('utf-8')
                logger.info(f"Read TXT file, content length: {len(content)}")

            elif filename_lower.endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(file)

                    # Extract content from paragraphs
                    content_parts = []

                    # CRITICAL: Extract header text first (this is where names/contact info often are)
                    for section in doc.sections:
                        # Get header content from first page header
                        header = section.header
                        if header:
                            for para in header.paragraphs:
                                if para.text.strip():
                                    content_parts.append(para.text)
                                    logger.info(f"📋 Header text found: {para.text[:100]}")

                    # Then add main document content
                    for para in doc.paragraphs:
                        if para.text.strip():
                            content_parts.append(para.text)

                    # Also check for text in tables (contact info is sometimes in tables)
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                            if row_text:
                                content_parts.append(row_text)

                    content = "\n".join(content_parts)
                    logger.info(f"Read DOCX file, content length: {len(content)}")
                    logger.info(f"📋 First 300 chars: {content[:300]}")
                except ImportError:
                    logger.error("DOCX support not installed")
                    return None, "❌ DOCX support not installed. Please install python-docx: pip install python-docx"
                except Exception as e:
                    logger.error(f"Error reading DOCX: {e}")
                    return None, f"❌ Error reading DOCX file: {str(e)}. The file may be corrupted."

            elif filename_lower.endswith('.pdf'):
                try:
                    import fitz  # PyMuPDF
                    pdf_bytes = file.read()
                    logger.info(f"Read PDF file, size: {len(pdf_bytes)} bytes")

                    if len(pdf_bytes) == 0:
                        logger.error("PDF file is empty")
                        return None, "❌ PDF file is empty. Please upload a valid PDF file."

                    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                        if len(doc) == 0:
                            logger.error("PDF has no pages")
                            return None, "❌ PDF file has no pages. Please upload a valid PDF file."

                        # Use "blocks" extraction which preserves layout better
                        # This helps capture header text that might be positioned at the top
                        content_parts = []
                        for page in doc:
                            # Extract text with layout preservation
                            # "blocks" mode maintains the order of text as it appears visually
                            text = page.get_text("text", sort=True)
                            if text.strip():
                                content_parts.append(text)

                        content = "\n".join(content_parts)
                        logger.info(f"Extracted PDF content, length: {len(content)}")
                        logger.info(f"📋 First 300 chars: {content[:300]}")
                except ImportError:
                    logger.error("PDF support not installed")
                    return None, "❌ PDF support not installed. Please install PyMuPDF: pip install PyMuPDF"
                except Exception as e:
                    logger.error(f"Error reading PDF: {e}")
                    return None, f"❌ Error reading PDF file: {str(e)}. The file may be corrupted or password-protected."
            else:
                # Try to read as text with error handling
                try:
                    content = file.read().decode('utf-8', errors='replace')
                    logger.info(f"Read as text file, content length: {len(content)}")
                except Exception as e:
                    logger.error(f"Error reading file as text: {e}")
                    return None, f"❌ Unsupported file format: {filename}. Please upload a TXT, PDF, or DOCX file."

            # Validate content was successfully extracted
            if content is None:
                logger.error("Content is None after processing")
                return None, f"❌ Could not extract text from {filename}. Please try a different file format."

            # Validate content
            if not content or len(content.strip()) < 50:
                logger.warning(f"Content too short: {len(content) if content else 0} characters")
                return None, f"❌ File appears to be empty or too short ({len(content) if content else 0} characters). Please upload a file with more content."

            # Limit content size for processing
            if len(content) > 10000:
                content = content[:10000] + "\n\n[Content truncated for processing]"
                logger.info("Content truncated to 10000 characters")

            # Basic validation - check if it looks like a resume
            if not any(keyword in content.lower() for keyword in ['experience', 'education', 'skills', 'employment', 'work', 'resume', 'cv', 'professional', 'career']):
                logger.warning("Content doesn't appear to be a resume")
                return content, "⚠️ Warning: This doesn't appear to be a resume, but I'll try to process it anyway."

            logger.info("File processing completed successfully")
            return content, None

        except UnicodeDecodeError as e:
            logger.error(f"Unicode decode error: {e}")
            return None, "❌ Could not read file - it may be corrupted or in an unsupported encoding format."
        except Exception as e:
            logger.error(f"Unexpected file processing error: {e}", exc_info=True)
            return None, f"❌ Unexpected error reading file: {str(e)}\n\nPlease try again or use a different file."
